import requests
from docx import Document
import os
from langchain_openai.chat_models.azure import AzureChatOpenAI
import requests
from bs4 import BeautifulSoup


####################################################################################################
## ESTE PROJETO FOI FEITO NO GOOGLE COLAB - EXECUTE O CÓDIGO LÁ PARA UM FUNCIONAMENTO MAIS RÁPIDO ##
####################################################################################################

subscription_key = "Executei o teste e removi a chave"
endpoint = 'http://api.cognitive.microsofttranslator.com'
location = "eastus2"
language_destination = 'pt-br'

def translate_text(text, target_laguage):
    path = '/translate'
    constructed_url = endpoint + path
    headers = {
        'Ocp-Apim-Subscription-Key': subscription_key,
        'Ocp-Apim-Subscription-Region': location,
        'Content-type': 'application/json',
        'X-ClientTraceId': str(os.urandom(16))
    }

    body = [{
        'text': text
    }]
    params = {
        'api-version': '3.0',
        'from': 'de-DE',
        'to': [language_destination]
    }
    request = requests.post(constructed_url, params=params, headers=headers, json=body)
    response = request.json()
    return response[0]["translations"][0]["text"]

translate_text("Eins, hier kommt die Sonne. Zwei, hier kommt die Sonne. Drei, sie ist der hellste Stern von allen Vier, hier kommt die Sonne", language_destination)

def translate_document(path):
  document = Document(path)
  full_text = []
  for paragraph in document.paragraphs:
    texto_traduzido = translate_text(paragraph.text, language_destination)
    full_text.append(texto_traduzido)

  documento_traduzido = Document()
  for line in full_text:
    documento_traduzido.add_paragraph(line)
  path_traduzido = path.replace('.docx', '_traduzido.docx')
  documento_traduzido.save(path_traduzido)
  return path_traduzido

input_file = 'content/teste.docx'
translate_document(input_file)


def extract_text_from_url(url):
    response = requests.get(url)

    if response.status_code != 200:
      soup = BeautifulSoup(response.text, 'html.parser')
      for script_or_style in soup(["script", "style"]):
          script_or_style.decompose()
      texto = soup.get_text(separator = ' ')
      #Limpar Texto
      linhas = (line.strip() for line in texto.splitlines())
      parts = (phrase.strip() for line in linhas for phrase in line.split("  "))
      texto_limpo = '\n'.join(chunk for chunk in parts if chunk)
      return texto_limpo

    else:
      print(f"Erro ao acessar a URL: {response.status_code}")
      return None

extract_text_from_url('https://dev.to/kenakamu/azure-open-ai-in-vnet-3alo')

client = AzureChatOpenAI(
    azure_endpoint = "https://deletei o recurso depois do teste",
    api_key = "deletei depois do teste",
    api_version = "2024-02-15-preview",
    deployment_name = "gpt-4o-mini",
    max_retries = 0
)

def translate_article(text, lang):
  messages = [
      ("system", "Você atua como um tradutor de textos"),
      ("user", f"Traduza o {text} para o idioma {lang} e responda em markdown")
  ]

  response = client.chat(messages)
  print(response.message.content)
  return response.message.content

translate_article("Hello World", "portuges")

url = 'https://dev.to/kenakamu/azure-open-ai-in-vnet-3alo'
text = extract_text_from_url(url)
article = translate_article(text, "pt-br")

